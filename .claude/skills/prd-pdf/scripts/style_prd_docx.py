#!/usr/bin/env python3
"""Apply the 8090 PRD house style to a pandoc-generated .docx, in place.

Modeled on packpilot/PRD/process/pipeline.py `format_docx` (the canonical 8090 PRD
formatter). One typeface (IBM Plex Mono), navy headings, Letter / 1in margins,
running header + page-number footer, content-proportional fixed table widths with a
navy header row, code blocks forced to 8pt mono so ASCII diagrams do not wrap,
natural section flow with keep-with-next (no break-before-every-heading), and a
single break that separates the cover/TOC front matter from the body.

Usage:
  style_prd_docx.py FILE.docx --title "..." --subtitle "..." \
      --version "Version 2.1 | 2026-06-30 | 8090.inc" --header "KYAI PRD v2.1"

Re-runnable: it replaces the owned header/footer rather than aborting.
"""
import sys, argparse, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY  = RGBColor(0x00, 0x33, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DG    = RGBColor(0x33, 0x33, 0x33)
MG    = RGBColor(0x55, 0x55, 0x55)
G6    = RGBColor(0x66, 0x66, 0x66)
FONT  = 'IBM Plex Mono'

PORTRAIT_W = 9360   # 6.5in in twips (text column at 1in margins on Letter)


def set_rfonts(rpr):
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), FONT)


def style_font_lock(style):
    rpr = style.element.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr'); style.element.append(rpr)
    set_rfonts(rpr)


def lock_run_font(run):
    run.font.name = FONT
    set_rfonts(run._element.get_or_add_rPr())


def add_page_field(paragraph):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin'); run._element.append(b)
    ir = OxmlElement('w:r'); it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '; ir.append(it); run._element.addnext(ir)
    sr = OxmlElement('w:r'); s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate'); sr.append(s); ir.addnext(sr)
    er = OxmlElement('w:r'); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); er.append(e); sr.addnext(er)
    lock_run_font(run)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx')
    ap.add_argument('--title', default='')
    ap.add_argument('--subtitle', default='Product Requirements Document')
    ap.add_argument('--version', default='')
    ap.add_argument('--header', default='PRD')
    a = ap.parse_args()

    doc = Document(a.docx)
    names = {s.name for s in doc.styles}
    # pandoc code-block / verbatim style names vary; match defensively
    code_styles = {n for n in names if any(k in n.lower() for k in ('source code', 'verbatim', 'sourcecode', 'preformatted'))}

    # 1. base styles: font, size, color, rFonts
    style_config = {
        'Normal': (10.5, None), 'Body Text': (10.5, None), 'First Paragraph': (10.5, None),
        'Heading 1': (20, NAVY), 'Heading 2': (16, NAVY), 'Heading 3': (12, DG), 'Heading 4': (11, MG),
    }
    for name, (size, color) in style_config.items():
        if name in names:
            st = doc.styles[name]
            st.font.name = FONT; st.font.size = Pt(size)
            if color is not None:
                st.font.color.rgb = color
            style_font_lock(st)
    for name, (before, after) in {'Heading 1': (36, 12), 'Heading 2': (24, 8), 'Heading 3': (14, 4), 'Heading 4': (10, 4)}.items():
        if name in names:
            pf = doc.styles[name].paragraph_format
            pf.space_before = Pt(before); pf.space_after = Pt(after); pf.keep_with_next = True
    for sn in ('Normal', 'Body Text', 'First Paragraph'):
        if sn in names:
            doc.styles[sn].paragraph_format.line_spacing = 1.08
    # code styles -> 8pt mono so 70-col ASCII fits the 6.5in column
    for cn in code_styles:
        st = doc.styles[cn]
        try:
            st.font.name = FONT; st.font.size = Pt(8); style_font_lock(st)
        except Exception:
            pass

    # 2. cover: the single H1 is the title; the subtitle/version lines follow
    for para in doc.paragraphs[:14]:
        t = para.text.strip()
        sn = para.style.name if para.style else ''
        if sn in ('Heading 1', 'Title'):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = NAVY
            para.paragraph_format.space_before = Pt(36); para.paragraph_format.space_after = Pt(18)
            para.paragraph_format.page_break_before = False
        elif a.subtitle and t == a.subtitle:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.font.size = Pt(14); r.font.color.rgb = G6
            para.paragraph_format.space_after = Pt(4)
        elif a.version and t == a.version:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = G6
            para.paragraph_format.space_after = Pt(14)

    # 3. images normalized to the 6.5in text column (scale down oversized; cap 8in tall)
    MAXW, MAXH = 5943600, 7315200
    for para in doc.paragraphs:
        for run in para.runs:
            for ext in run._element.findall('.//' + qn('wp:extent')):
                cx, cy = int(ext.get('cx', 0)), int(ext.get('cy', 0))
                if not cx or not cy:
                    continue
                scale = min(1.0, MAXW / cx)
                if cy * scale > MAXH:
                    scale = MAXH / cy
                if abs(scale - 1.0) > 0.01:
                    ncx, ncy = int(cx * scale), int(cy * scale)
                    ext.set('cx', str(ncx)); ext.set('cy', str(ncy))
                    for ae in run._element.findall('.//' + qn('a:ext')):
                        if int(ae.get('cx', 0)) == cx:
                            ae.set('cx', str(ncx)); ae.set('cy', str(ncy))

    # 4. tables: fixed layout, content-proportional widths, navy header row, no landscape
    for table in doc.tables:
        ncols = len(table.columns)
        if ncols == 0:
            continue
        rows = table.rows
        body_size = 8 if ncols >= 5 else 9
        # content-proportional widths with an 8% floor, summing to the text column
        maxlen = []
        for c in range(ncols):
            ml = 1
            for r in range(len(rows)):
                try:
                    ml = max(ml, len(rows[r].cells[c].text))
                except Exception:
                    pass
            maxlen.append(min(ml, 120))
        floor = int(PORTRAIT_W * 0.08)
        tot = sum(maxlen) or 1
        raw = [max(floor, int(PORTRAIT_W * m / tot)) for m in maxlen]
        s = sum(raw) or 1
        widths = [max(1, int(w * PORTRAIT_W / s)) for w in raw]
        # tblPr: fixed layout, fixed total width
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
        for tag, attrs in (('w:tblW', {'w:w': str(PORTRAIT_W), 'w:type': 'dxa'}),
                           ('w:tblLayout', {'w:type': 'fixed'})):
            el = tblPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag); tblPr.append(el)
            for k, v in attrs.items():
                el.set(qn(k), v)
        # grid
        grid = tbl.find(qn('w:tblGrid'))
        if grid is not None:
            tbl.remove(grid)
        grid = OxmlElement('w:tblGrid')
        for w in widths:
            gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); grid.append(gc)
        tblPr.addnext(grid)
        # cells: width + font
        for row in rows:
            for ci, cell in enumerate(row.cells):
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr'); cell._tc.insert(0, tcPr)
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
                tcW.set(qn('w:w'), str(widths[ci] if ci < len(widths) else widths[-1]))
                tcW.set(qn('w:type'), 'dxa')
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = FONT; r.font.size = Pt(body_size); lock_run_font(r)
        # keep rows whole across page breaks; repeat the header row on each page
        for ri, row in enumerate(rows):
            trPr = row._tr.get_or_add_trPr()
            trPr.append(OxmlElement('w:cantSplit'))
            if ri == 0:
                trPr.append(OxmlElement('w:tblHeader'))
        # header row: navy fill, white bold
        for cell in rows[0].cells:
            tcPr = cell._tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr'); cell._tc.insert(0, tcPr)
            old = tcPr.find(qn('w:shd'))
            if old is not None:
                tcPr.remove(old)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '003366')
            tcPr.append(shd)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = WHITE; r.font.bold = True; r.font.size = Pt(9)

    # 5. page setup: Letter, 1in margins
    for section in doc.sections:
        sectPr = section._sectPr
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is not None:
            pgSz.set(qn('w:w'), '12240'); pgSz.set(qn('w:h'), '15840')
            if pgSz.get(qn('w:orient')):
                del pgSz.attrib[qn('w:orient')]
        pgMar = sectPr.find(qn('w:pgMar'))
        if pgMar is not None:
            for s in ('w:top', 'w:right', 'w:bottom', 'w:left'):
                pgMar.set(qn(s), '1440')

    # 6. running header + page-number footer (replace owned content; re-runnable)
    for section in doc.sections:
        header = section.header; header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.clear(); hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = hp.add_run(a.header); run.font.name = FONT; run.font.size = Pt(8); run.font.color.rgb = NAVY
        lock_run_font(run)
        footer = section.footer; footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]; fp.clear(); fp.add_run('\t')
        add_page_field(fp)
        pPr = fp._element.get_or_add_pPr()
        tabs = OxmlElement('w:tabs'); tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right'); tab.set(qn('w:pos'), str(PORTRAIT_W)); tab.set(qn('w:leader'), 'none')
        tabs.append(tab); pPr.append(tabs)

    # 7. break that separates front matter (cover + TOC) from the body: the FIRST numbered H2
    for para in doc.paragraphs:
        if para.style and para.style.name == 'Heading 2' and re.match(r'^\d+\.', para.text.strip()):
            para.paragraph_format.page_break_before = True
            break

    # 8. final font-lock pass over every run (body + tables + headers/footers + TOC)
    def lock_para(p):
        for r in p.runs:
            lock_run_font(r)
    for p in doc.paragraphs:
        lock_para(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    lock_para(p)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                lock_para(p)

    # 9. force field update on load so LibreOffice fills the TOC page numbers on --convert-to pdf
    try:
        settings = doc.settings.element
        if settings.find(qn('w:updateFields')) is None:
            uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true'); settings.append(uf)
    except Exception:
        pass

    doc.save(a.docx)
    print(f"styled: {a.docx}  ({len(doc.tables)} tables, code styles: {sorted(code_styles) or 'none'})")


if __name__ == '__main__':
    main()
